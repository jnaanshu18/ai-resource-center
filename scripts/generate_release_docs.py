#!/usr/bin/env python3
"""Generate user-friendly Word release guides from markdown sources."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "release"
SITE_URL = "https://ai.dailycodesolutions.com/"
UPDATED = "26 August 2026"

# Brand-ish colors
HEADER_FILL = "1F4E79"  # dark blue
ALT_ROW_FILL = "F2F7FB"
PLACEHOLDER_FILL = "E8EEF4"


def set_cell_shading(cell, fill_hex: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_table(table, header_fill: str = HEADER_FILL) -> None:
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(table.rows[1:], start=1):
        if i % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, ALT_ROW_FILL)


def add_title(doc: Document, title: str, subtitle: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    if subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(subtitle)
        sr.font.size = Pt(11)
        sr.font.color.rgb = RGBColor(89, 89, 89)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Site: {SITE_URL}\nLast updated: {UPDATED}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(89, 89, 89)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    style_table(table)
    doc.add_paragraph()


def add_screenshot_placeholder(doc: Document, caption: str) -> None:
    """Reserved box for a screenshot the author can paste over later."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Screenshot: {caption} ]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Gray bordered paragraph as visual placeholder (approx 4:3 box)
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = box.add_run("\n\n\n\n\n")
    br.font.size = Pt(8)

    # Light shading via paragraph border trick — use a one-cell table instead
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_shading(cell, PLACEHOLDER_FILL)
    cell.text = ""
    inner = cell.paragraphs[0]
    inner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ph = inner.add_run("Paste screenshot here")
    ph.italic = True
    ph.font.size = Pt(10)
    ph.font.color.rgb = RGBColor(120, 120, 120)
    cell.width = Inches(5.5)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.size = Pt(9)
    cr.font.color.rgb = RGBColor(89, 89, 89)
    doc.add_paragraph()


def add_checklist_table(doc: Document, section: str, items: list[str]) -> None:
    add_heading(doc, section, level=2)
    table = doc.add_table(rows=1 + len(items), cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Done"
    table.rows[0].cells[1].text = "Task"
    style_table(table)
    for i, item in enumerate(items, start=1):
        table.rows[i].cells[0].text = "☐"
        table.rows[i].cells[1].text = item
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def build_employee_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title(doc, "DCS AI Resource Center", "Employee Guide")
    add_body(
        doc,
        "Welcome! This guide shows you how to use our internal AI tools site — "
        "where to go, what each section does, and how to suggest tools or share wins.",
    )

    add_heading(doc, "What is the AI Resource Center?")
    add_body(
        doc,
        "The DCS AI Resource Center is our internal website for finding AI tools "
        "the team already trusts and learning how others use them.",
    )
    add_bullets(
        doc,
        [
            "Browse approved AI tools (Directory)",
            "Get pointed to the right tool for a job (AI hub + Choose a tool)",
            "Copy ready-made prompts",
            "Read team stories and share your own wins",
            "Suggest new tools for review (Suggestions)",
        ],
    )

    add_heading(doc, "How to sign in")
    add_numbered(
        doc,
        [
            f"Open {SITE_URL} in your browser.",
            "Sign in with the team username and password shared internally by your admin "
            "(or use an invite link if you were given one).",
            "After sign-in you land on AI hub.",
        ],
    )
    add_body(
        doc,
        "Important: Do not share login credentials in public channels "
        "(Slack, email to external parties, etc.).",
        bold=True,
    )
    add_screenshot_placeholder(doc, "Login screen")

    add_heading(doc, "Where to go — site map")
    add_table(
        doc,
        ["Tab", "What it is", "When to use it"],
        [
            ["AI hub", "Job chooser, Start here, Tool of the week", "I have a task — point me to trusted tools"],
            ["Directory", "Full catalog of approved tools", "Search, filter, compare, open tool detail"],
            ["Choose a tool", "Category guides + tool-vs-tool matchups", "Which option in this category?"],
            ["Prompts", "Searchable prompt library", "Copy prompts for ChatGPT, Claude, Cursor, etc."],
            ["Team stories", "Real wins + learning links", "See how teammates use tools; share your story"],
            ["Suggestions", "Tool suggestion queue", "Propose a tool the team should evaluate"],
        ],
    )
    add_body(
        doc,
        "Each sign-in: AI hub shows a short orientation banner. Click Got it to dismiss for that visit — it returns on your next sign-in.",
    )
    add_body(doc, "The logo / company name in the header always takes you back to AI hub.")
    add_screenshot_placeholder(doc, "AI hub — main navigation and orientation banner")

    add_heading(doc, "Directory — what’s in it")
    add_body(doc, "Only approved tools appear in the Directory:")
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["Production", "Core daily tools (e.g. Cursor, OpenClaw)"],
            ["Approved", "Cleared for team use"],
            ["Archived", "Kept for reference only"],
        ],
    )
    add_body(
        doc,
        "Tools still being evaluated appear under Suggestions (New / In review / Rejected) — "
        "not in the Directory until an admin approves them.",
    )
    add_screenshot_placeholder(doc, "Directory — search, filters, and tool cards")

    add_heading(doc, "Tool detail pages")
    add_body(doc, "When you open a tool from the Directory, you will see:")
    add_bullets(
        doc,
        [
            "Tutorial video and key facts",
            "When to use, alternatives, and team notes",
            "Limitations, cost & security, and DCS evaluation (when available)",
            "Related tool-vs-tool links, team stories, and prompts",
            "Visit website link next to the tool name",
            "Add to compare — pick up to 3 tools for side-by-side comparison",
        ],
    )
    add_screenshot_placeholder(doc, "Tool detail page — overview and related links")

    add_heading(doc, "How to contribute")
    add_heading(doc, "Suggest a tool", level=2)
    add_numbered(
        doc,
        [
            "Go to Suggestions → suggest form (or Contribute tab).",
            "Enter tool name, website URL (required), and a short note.",
            "Your suggestion appears in the queue as New — an admin will review it.",
        ],
    )
    add_screenshot_placeholder(doc, "Suggestions — suggest a tool form")

    add_heading(doc, "Share a team win", level=2)
    add_numbered(
        doc,
        [
            "Go to Team stories → Share a win form.",
            "Fill in title, tool, impact (max 400 characters), and how you did it.",
            "After admin approval, your story appears on Team stories.",
        ],
    )
    add_screenshot_placeholder(doc, "Team stories — share a win form")

    add_heading(doc, "Data safety")
    add_bullets(
        doc,
        [
            "Use only approved tools from the Directory for work-related tasks unless your lead says otherwise.",
            "Never paste client secrets, personal data (PII), or production credentials into forms or unapproved AI tools.",
            "All directory tools are tagged Internal — treat outputs and uploads accordingly.",
        ],
    )

    add_heading(doc, "Need help?")
    add_table(
        doc,
        ["Issue", "What to do"],
        [
            ["Can’t log in", "Contact your site admin (credentials may have been rotated)"],
            ["Tool missing from Directory", "It may still be under review — check Suggestions, or suggest it"],
            ["Suggestion stuck", "Ask an admin to review it on the Suggestions tab"],
            ["Wrong info on a tool", "Tell an admin — they can update the catalog"],
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Daily Code Solutions — Internal use only")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def build_admin_doc() -> Document:
    """Admin guide — review workflow only, no maintainer/technical content."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title(doc, "DCS AI Resource Center", "Admin Guide")
    add_body(
        doc,
        "For admins who suggest tools, review the queue, and approve team wins. "
        "Share the Employee Guide with the wider team — not this document.",
        bold=True,
    )

    add_heading(doc, "For v1 — main purpose for everyone (including admins)")
    add_body(
        doc,
        "The number-one job at launch is to suggest AI tools the team should evaluate. "
        "That applies to you as an admin too — not only employees. Browse what you use or hear about, "
        "and add it to the queue. The more quality suggestions we collect in v1, the stronger the Directory becomes.",
    )
    add_body(
        doc,
        "You can also review what others submit and approve team wins — but suggesting tools comes first.",
        bold=True,
    )

    add_heading(doc, "How to suggest a tool (start here)")
    add_numbered(
        doc,
        [
            f"Open {SITE_URL} and sign in with admin@dailycodesolutions.com.",
            "Go to Suggestions → suggest form (or Contribute tab).",
            "Enter tool name, website URL (required), and a short note on why the team should look at it.",
            "Submit — your suggestion appears in the queue as New (same as any employee submission).",
            "You or another admin can Review it later (see below).",
        ],
    )
    add_body(
        doc,
        "Tip: Suggest tools you personally use or want the team to try — don’t wait for someone else to add them.",
    )
    add_body(
        doc,
        "Data safety: Never paste client secrets, personal data, or production credentials into the form.",
        bold=True,
    )
    add_screenshot_placeholder(doc, "Suggestions — suggest a tool form")

    add_heading(doc, "Sign in")
    add_numbered(
        doc,
        [
            f"Open {SITE_URL}",
            "Sign in with admin@dailycodesolutions.com and the admin password (shared securely with admins only).",
            "You land on AI hub — use the top navigation like any employee.",
        ],
    )
    add_body(doc, "Keep admin credentials private. Do not post them in public Slack channels or email threads.", bold=True)
    add_screenshot_placeholder(doc, "Admin sign-in — header shows Admin badge")

    add_heading(doc, "Your role")
    add_table(
        doc,
        ["Role", "Login", "What they do"],
        [
            ["Employee", "Shared team account or invite link", "Browse the site, suggest tools, share wins"],
            [
                "Admin",
                "admin@dailycodesolutions.com",
                "Suggest tools like everyone else, plus review the queue and approve team wins",
            ],
        ],
    )
    add_body(doc, "After sign-in, your header badge shows Admin (employees see Team).")

    add_heading(doc, "Review tool suggestions")
    add_body(doc, "Go to Suggestions. Each card has a Review button (admin only).")
    add_screenshot_placeholder(doc, "Suggestions queue — Review button on a card")

    add_heading(doc, "Step 1 — Open Review", level=2)
    add_body(doc, "Click Review on a suggestion. You will see the tool name, link, and any notes.")

    add_heading(doc, "Step 2 — Choose an action", level=2)
    add_table(
        doc,
        ["Action", "When to use it"],
        [
            [
                "Assign",
                "Hand the suggestion to someone to evaluate. Pick a name and optional comment. Status becomes In review.",
            ],
            ["Reject", "The tool is not a fit. You must add a comment explaining why. Status becomes Rejected."],
            [
                "Add to Directory",
                "The tool is cleared for the team. Only when status is In review. Opens a short catalog form.",
            ],
        ],
    )
    add_screenshot_placeholder(doc, "Review panel — Assign, Reject, Add to Directory")

    add_heading(doc, "Step 3 — Add to Directory (when approving)", level=2)
    add_numbered(
        doc,
        [
            "Fill required fields: category, pricing, description, tutorial video URL.",
            "Click Approve & queue publish.",
            "The site marks the suggestion Approved and downloads a JSON file to your computer.",
            "Send that JSON file to the site maintainer — they publish it to the live Directory.",
        ],
    )
    add_body(
        doc,
        "The tool will not appear in Directory until the maintainer publishes it (usually within a day).",
        bold=True,
    )
    add_screenshot_placeholder(doc, "Add to Directory form and success screen with downloaded file")

    add_heading(doc, "Suggestion statuses", level=2)
    add_table(
        doc,
        ["Status", "Meaning"],
        [
            ["New", "Just submitted — assign someone or reject"],
            ["In review", "Someone is evaluating — approve for Directory or reject"],
            ["Approved", "Cleared — waiting for maintainer to publish live"],
            ["Rejected", "Not approved — reason is in the admin note"],
        ],
    )

    add_heading(doc, "Approve team wins")
    add_body(doc, "When someone uses Team stories → Share a win, their story goes to a Google Sheet (Team wins tab).")
    add_numbered(
        doc,
        [
            "Open the Team wins sheet (link from site maintainer).",
            "Find the new row and read the submission — check for client secrets or personal data.",
            "Set Status to Approved.",
        ],
    )
    add_body(doc, "Approved wins appear on Team stories automatically. Contact the submitter if it needs edits.")
    add_screenshot_placeholder(doc, "Google Sheet — Team wins tab, Status column")

    add_heading(doc, "Share access with employees")
    add_table(
        doc,
        ["Task", "What to do"],
        [
            ["Give employees site access", "Share the team username and password securely"],
            ["Invite link", "Ask the site maintainer to generate one"],
            ["Rotate team password", "Ask the site maintainer"],
        ],
    )

    add_heading(doc, "Remind the team")
    add_bullets(
        doc,
        [
            "Directory = tools DCS has already approved for use.",
            "Suggestions = tools still being evaluated — not approved yet.",
            "Never paste client data, passwords, or production secrets into AI tools or forms.",
        ],
    )

    add_heading(doc, "Employee announcement (copy-paste draft)")
    quote = doc.add_paragraph()
    for line in [
        "Subject: DCS AI Resource Center is live",
        "",
        "We’ve launched an internal site to help you find AI tools DCS already trusts:",
        SITE_URL,
        "",
        "• Directory — approved tools with tutorials and compare",
        "• Choose a tool — guides when you’re deciding between options",
        "• Prompts — copy-ready prompts for common tasks",
        "• Team stories — real examples from the team",
        "• Suggestions — propose tools for us to review (main ask for v1)",
        "",
        "Sign in with the shared team credentials we’ve sent separately.",
        "Questions? Reply in [your channel] or contact [admin name].",
    ]:
        r = quote.add_run(line + "\n")
        r.font.size = Pt(10)
        r.italic = True

    add_heading(doc, "When to contact the site maintainer")
    add_table(
        doc,
        ["Situation", "They can help with"],
        [
            ["Tool approved but not in Directory yet", "Publish the JSON file you downloaded"],
            ["Site won’t load or looks broken", "Server / deploy issues"],
            ["Login stopped working for everyone", "Password or config update"],
            ["Wrong text on a tool page", "Update catalog data"],
            ["Need a new admin password", "Rotate admin credentials"],
            ["Sheet not receiving submissions", "Backend integration"],
        ],
    )

    add_heading(doc, "What’s on the site today (v1)")
    add_body(
        doc,
        "15 directory tools · 16 prompts · 10 team stories · 9 learning links · "
        "7 comparisons · 10 AI hub jobs",
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Daily Code Solutions — Admin use only")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def build_maintainer_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_title(doc, "DCS AI Resource Center", "Maintainer Guide (v1.0)")
    add_body(
        doc,
        "For site maintainer(s) only — deploy, config, CSV/git, scripts. "
        "Share the Admin Guide with reviewers; share the Employee Guide with the team.",
        bold=True,
    )

    add_heading(doc, "Roles")
    add_table(
        doc,
        ["Role", "Login", "Can do"],
        [
            ["Employee", "Shared team account or invite link", "Browse, suggest tools, share wins, compare"],
            [
                "Admin",
                "admin@dailycodesolutions.com",
                "Review suggestions, approve wins in Sheet, queue Directory publish (JSON handoff to you)",
            ],
            ["Maintainer", "(this guide)", "Deploy, config secrets, publish Directory tools, fix integrations"],
        ],
    )
    add_body(
        doc,
        "Login is a soft gate. Production must use Cloudflare Access (or equivalent) for network-level protection.",
    )

    add_heading(doc, "Publishing a tool to Directory")
    add_body(doc, "Admins approve in the UI and download a JSON file. You publish it:")
    code = doc.add_paragraph()
    code.style = "No Spacing"
    for line in [
        "python scripts/add_directory_tool.py --json directory-tool-name.json",
        "python scripts/generate_site_data.py",
        "git add data/ docs/data.js",
        'git commit -m "Add [Tool name] to directory"',
        "git push",
    ]:
        r = code.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
    add_body(
        doc,
        "Or edit data/ai_tools_directory.csv directly and push — GitHub Actions regenerates docs/data.js.",
    )

    add_heading(doc, "Launch checklist")
    add_body(doc, "Complete before announcing to the company.")

    add_checklist_table(
        doc,
        "Infrastructure",
        [
            "DNS for ai.dailycodesolutions.com points to GitHub Pages",
            "GitHub Settings → Pages → Source = GitHub Actions",
            "Secret SITE_CONFIG_JS set with production config",
            'environment: "production" in deployed config',
            "Cloudflare Access (or VPN) in front of production URL",
            "Push latest main and confirm Actions deploy succeeds",
        ],
    )
    add_checklist_table(
        doc,
        "Google Sheets / Apps Script",
        [
            "DCS Apps Script deployed as web app (Anyone access)",
            "submitUrl and winSubmit.submitUrl set in production config",
            "ASSIGN_SECRET Script property matches assignSecret in config",
            "Submissions + Team wins tabs receiving test rows",
            "Published CSV URLs work (or listWins action tested)",
        ],
    )
    add_checklist_table(
        doc,
        "Auth",
        [
            "Admin password hash updated and shared securely with admins only",
            "Employee (team) password shared securely with company — not in repo",
            'Invite token generated if using invite links: python scripts/hash_auth_secret.py "token"',
            "Test login on production in incognito (admin + team)",
        ],
    )
    add_checklist_table(
        doc,
        "Content",
        [
            "data/ai_tools_directory.csv reviewed — descriptions, tutorials, statuses",
            "data/tool_submissions.csv queue reflects current trial tools",
            "data/team_use_cases.csv + approved sheet wins reviewed",
            "data/site_highlights.csv — Start here + Tool of the week current",
            "Run python scripts/generate_site_data.py and commit if CSV changed locally",
        ],
    )
    add_checklist_table(
        doc,
        "Quality",
        [
            "python scripts/qa_full_site.py — target 0 FAIL",
            "node --check docs/script.js",
            "Smoke test all 6 nav tabs on production",
            "Test suggest-a-tool + share-a-win end-to-end",
            "Test tool compare (pick 2 tools → compare strip → side-by-side)",
        ],
    )
    add_checklist_table(
        doc,
        "Communication",
        [
            "Announce URL + team login to employees",
            "Share Employee Guide (Word/PDF) — not Maintainer guide",
            "Brief admins on Admin Guide review workflow",
        ],
    )

    add_heading(doc, "Deploy process")
    add_table(
        doc,
        ["Step", "Action"],
        [
            ["1", "Edit data/*.csv (or docs/*.js / css for UI changes)"],
            ["2", "Run python scripts/generate_site_data.py (only after CSV edits)"],
            ["3", "git commit + push to main"],
            ["4", "GitHub Actions regenerates data.js and deploys Pages (~2–3 min)"],
            ["5", "Optional: python scripts/qa_full_site.py"],
        ],
    )
    add_body(
        doc,
        "Do not commit docs/site-config.js — production config lives in GitHub secret SITE_CONFIG_JS.",
        bold=True,
    )

    add_heading(doc, "v1 scope — intentionally off")
    add_table(
        doc,
        ["Feature", "Config / code", "Enable when"],
        [
            ["Tool assignment UI", "toolAssignments.enabled: false", "Ready to assign Testing/Exploring tools"],
            ["Assignment email notify", "toolAssignments.notify.enabled: false", "Worker deployed + Resend configured"],
            ["Team self-registration", "teamDirectory.allowSelfRegister: false", "Want open signup form"],
            ["Client-work + internal contact", "TOOL_TRUST_ROW_ENABLED = false", "Policy/contact data ready"],
            ["Legacy full Add a tool form", "contribute.fullForm.enabled: false", "Not planned for v1"],
        ],
    )

    add_heading(doc, "QA status (26 Aug 2026)")
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["PASS", "190"],
            ["FAIL", "1"],
            ["WARN", "1"],
            ["TOTAL", "192"],
        ],
    )
    add_table(
        doc,
        ["Result", "Item"],
        [
            ["FAIL", "Comparisons missing scenario field — cosmetic; fix in data/tool_comparison.csv if desired"],
            ["WARN", "No assignees in CSV (expected — assignments off for v1)"],
        ],
    )
    add_body(doc, "All critical paths pass: auth, directory, submissions, team stories, prompts, guides, compare, deploy.")

    add_heading(doc, "Security and compliance")
    add_bullets(
        doc,
        [
            "Passwords stored as SHA-256 hashes only — never plaintext in git.",
            "docs/site-config.js and docs/site-config.production.js are gitignored.",
            "Site config injected at deploy via SITE_CONFIG_JS secret.",
            "All directory tools tagged Internal — remind team about client-confidential content.",
            "Submitter names hidden on public Suggestions list.",
            "Login rate limit: 5 attempts → 60s lockout.",
            "Rotate employee password: python scripts/hash_auth_secret.py, update SITE_CONFIG_JS, redeploy.",
        ],
    )

    add_heading(doc, "Support runbook")
    add_table(
        doc,
        ["Topic", "Action"],
        [
            ["Site down / deploy failed", "GitHub Actions → re-run workflow"],
            ["Wrong tool data", "Edit CSV → push"],
            ["Admin JSON waiting to publish", "add_directory_tool.py --json → push"],
            ["Login issues", "Rotate hash, update SITE_CONFIG_JS, redeploy"],
            ["Apps Script errors", "Sheet + Apps Script execution log"],
        ],
    )

    add_heading(doc, "Useful commands", level=2)
    cmd = doc.add_paragraph()
    for line in [
        "python scripts/generate_site_data.py",
        "python scripts/qa_full_site.py",
        'python scripts/hash_auth_secret.py "new-password"',
        'python scripts/add_team_member.py --name "Name" --email "name@dailycodesolutions.com"',
        "python scripts/add_directory_tool.py --help",
        "node --check docs/script.js",
    ]:
        r = cmd.add_run(line + "\n")
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Daily Code Solutions — Maintainer use only")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    employee_path = OUT_DIR / "DCS-AI-Resource-Center-Employee-Guide.docx"
    admin_path = OUT_DIR / "DCS-AI-Resource-Center-Admin-Guide.docx"
    maintainer_path = OUT_DIR / "DCS-AI-Resource-Center-Maintainer-Guide.docx"

    for path, builder in [
        (employee_path, build_employee_doc),
        (admin_path, build_admin_doc),
        (maintainer_path, build_maintainer_doc),
    ]:
        try:
            builder().save(path)
            print(f"Created: {path}")
        except PermissionError:
            alt = path.with_name(path.stem + "-updated.docx")
            builder().save(alt)
            print(f"Created (file locked): {alt}")
    print("\nTip: Open in Word, paste screenshots into the gray placeholder boxes,")
    print("     then File > Save As > PDF when ready to share a read-only copy.")


if __name__ == "__main__":
    main()
